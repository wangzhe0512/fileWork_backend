#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""检查最新生成的报告质量"""

from docx import Document
import re

REPORT_NEW = r'd:\wangzhe\0.Qoder\fileWork_backend\docs\测试反馈\tp测试报告例子\spx测试\2024年度报告-19.docx'
TEMPLATE = r"d:\wangzhe\0.Qoder\fileWork_backend\docs\测试反馈\tp测试报告例子\spx测试\2024年子模板 (5).docx"
REPORT_OLD = r"d:\wangzhe\0.Qoder\fileWork_backend\docs\测试反馈\tp测试报告例子\spx测试\斯必克流体技术_FY2024 转让定价同期资料【讨论稿】_20250710 - 副本.docx"

def extract_placeholders(path):
    doc = Document(path)
    phs = set()
    pattern = r'\{\{([^}]+)\}\}'
    for para in doc.paragraphs:
        phs.update(re.findall(pattern, para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                phs.update(re.findall(pattern, cell.text))
    return phs

print("=" * 60)
print("📊 报告质量检查")
print("=" * 60)

# 1. 子模板占位符
print("\n【1】子模板占位符检查")
tpl_phs = extract_placeholders(TEMPLATE)
normal = [p for p in tpl_phs if not p.startswith('_')]
print(f"  总计: {len(tpl_phs)} 个占位符")
print(f"  普通占位符: {len(normal)} 个")
if normal:
    print("  未替换列表:")
    for p in sorted(normal)[:15]:
        print(f"    - {{{{ {p} }}}}")

# 2. 新报告占位符
print("\n【2】新报告占位符检查")
new_phs = extract_placeholders(REPORT_NEW)
new_normal = [p for p in new_phs if not p.startswith('_')]
print(f"  总计: {len(new_phs)} 个占位符")
if new_normal:
    print("  ⚠️ 未替换的普通占位符:")
    for p in sorted(new_normal):
        print(f"    - {{{{ {p} }}}}")
else:
    print("  ✅ 所有普通占位符已替换")

# 3. 基本对比
print("\n【3】基本对比")
old_doc = Document(REPORT_OLD)
new_doc = Document(REPORT_NEW)
print(f"  历史报告: {len(old_doc.tables)} 表格, {len(old_doc.paragraphs)} 段落")
print(f"  新报告: {len(new_doc.tables)} 表格, {len(new_doc.paragraphs)} 段落")

print("\n" + "=" * 60)
if new_normal:
    print("⚠️ 新报告存在未替换占位符")
else:
    print("✅ 检查通过")
